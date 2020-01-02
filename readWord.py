# 需要事先安装python-docx库
import json
import os
from docx import Document
import re

# 读取段落文字
def readPara(doc):  # doc为待读取的文书路径
    document = Document(doc)
    paras=list()
    data=list()
    dang=""
    for p in document.paragraphs:
        if p.text.strip() !="" :
            paras.append(p.text)
            # paras.insert(0,zhuti)
    for i in range(len(paras)):
        if '原告'not in paras[i] and '被告' not in paras[i]:
            data.append(paras[i])
        else:
            paras=paras[i:len(paras)]
            break
    p1 = re.compile('原告.*案')
    # p2=re.compile('被告.*：')
    for i in range(len(paras)):
        if len(p1.findall(paras[i]))==0:
            paras[i]=paras[i].replace("\n", "")
            dang=dang+paras[i].strip()
        else:
            paras=paras[i:len(paras)]
            break
    dang2 = re.split('被告', dang[10:], 1)
    dang2[0] = dang[0:10] + dang2[0]
    data.extend(dang2)
    data.extend(paras)
    return data

# 读取表格内容
def readTable(doc):
    content = []
    document = Document(doc)
    tables = document.tables
    for t in tables:
        for r in t.rows:
            for c in r.cells:
                text = ''
                for p in c.paragraphs:
                    text += p.text
                content.append(text)
    return content

#将doc处理成字典格式
def filter(root):
    # error=[]
    notread=[]
    data=[]
    p1 = re.compile('(.*)事实.*?理由(.*)')
    for f in os.listdir(root):
        root2=os.path.join(root, f)
        for key in os.listdir(root2) :
            key = os.path.join(root2,key)
            datatemp = dict()
            try:
                if key.endswith(".docx"):
                    paras = readPara(key)
                    if len(paras) == 0:
                        if ('民事判决书' in key or '民 事 判 决' in key):
                            notread.append(key)
                    if len(paras) > 0 and (paras[1] == '民 事 判 决 书' or paras[1] == '民事判决书'):
                        datatemp['key'] = key
                        datatemp['原告'] = paras[3]
                        datatemp['被告'] = paras[4]
                        datatemp['案由+程序'] = paras[5]
                        if '请求' in paras[6]:
                            temp = p1.findall(paras[6])
                            if len(temp) > 0:
                                temp = temp[0]
                                if len(temp[1]) > 5:
                                    datatemp['诉讼请求'] = temp[0]
                                    datatemp['事实理由'] = temp[1]
                                else:
                                    datatemp['诉讼请求'] = paras[6]
                                    datatemp['事实理由'] = paras[7]
                            else:
                                datatemp['诉讼请求'] = paras[6]
                                datatemp['事实理由'] = paras[7]
                            data.append(datatemp)
                        else:
                            notread.append(key)
            except:
                if ('民事判决书' in key or '民 事 判 决' in key):
                    notread.append(key)
                # error.append(key)
        #记录读取的内容
    with open("data.json", 'w', encoding='utf-8')as f2:
        f2.truncate()  # 清空文件内容
        for key in data:
            json.dump(key, f2, ensure_ascii=False)
            f2.write('\n')

    #记录未读的路径
    with open('recoder.txt', "w") as fw:
        fw.truncate()
        for i in range(len(notread)):
            print(str(notread[i]) , file=fw)

def filjson(root):
    fw=open('recoder.txt', "a")
    with open("data.json", 'r', encoding='utf-8')as f:
        data = []
        i=289
        for line in f:
            d = json.loads(line)  # 转成字典格式
            if len(d['事实理由'])<15:
                i=i+1
                print(str(d['key']), file=fw)

    fw.close()

if __name__ == '__main__':
    # root='temp'
    # filter(root)
    root='data.json'
    filjson(root)





